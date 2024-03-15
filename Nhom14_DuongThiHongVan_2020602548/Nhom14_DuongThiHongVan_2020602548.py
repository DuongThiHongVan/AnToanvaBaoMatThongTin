
from math import e
import random
from PyQt5 import QtCore, QtGui, QtWidgets
from PyQt5.QtWidgets import ( QFileDialog)
import docx
import pyperclip


class Ui_MainWindow(QtWidgets.QMainWindow):
    def setupUi(self, MainWindow):
        MainWindow.setObjectName("MainWindow")
        MainWindow.setEnabled(True)
        # cả khung
        MainWindow.resize(1111, 750)
        self.centralwidget = QtWidgets.QWidget(MainWindow)
        self.centralwidget.setObjectName("centralwidget")
        self.btnGKey = QtWidgets.QPushButton(self.centralwidget)
        # Khung chữ generation key
        self.btnGKey.setGeometry(QtCore.QRect(40, 20, 161, 61))
        font = QtGui.QFont()
        font.setPointSize(10)
        self.btnGKey.setFont(font)
        self.btnGKey.setObjectName("btnGKey")
        self.txtAlpha = QtWidgets.QLineEdit(self.centralwidget)
        self.txtAlpha.setGeometry(QtCore.QRect(230, 60, 71, 22))
        font = QtGui.QFont()
        font.setPointSize(10)
        self.txtAlpha.setFont(font)
        self.txtAlpha.setReadOnly(False)
        self.txtAlpha.setObjectName("txtAlpha")

        self.label = QtWidgets.QLabel(self.centralwidget)
        self.label.setGeometry(QtCore.QRect(230, 30, 55, 16))
        font = QtGui.QFont()
        font.setPointSize(10)
        self.label.setFont(font)
        self.label.setObjectName("label")

        self.label_2 = QtWidgets.QLabel(self.centralwidget)
        self.label_2.setGeometry(QtCore.QRect(320, 30, 55, 16))
        font = QtGui.QFont()
        font.setPointSize(10)
        self.label_2.setFont(font)
        self.label_2.setObjectName("label_2")

        self.txtBeta = QtWidgets.QLineEdit(self.centralwidget)
        self.txtBeta.setGeometry(QtCore.QRect(320, 60, 71, 22))
        font = QtGui.QFont()
        font.setPointSize(10)
        self.txtBeta.setFont(font)
        self.txtBeta.setReadOnly(True)
        self.txtBeta.setObjectName("txtBeta")

        self.label_3 = QtWidgets.QLabel(self.centralwidget)
        self.label_3.setGeometry(QtCore.QRect(410, 30, 55, 16))
        font = QtGui.QFont()
        font.setPointSize(10)
        self.label_3.setFont(font)
        self.label_3.setObjectName("label_3")

        self.txtP = QtWidgets.QLineEdit(self.centralwidget)
        self.txtP.setGeometry(QtCore.QRect(410, 60, 71, 22))
        font = QtGui.QFont()
        font.setPointSize(10)
        self.txtP.setFont(font)
        self.txtP.setObjectName("txtP")

        self.btnEncrypt = QtWidgets.QPushButton(self.centralwidget)
        self.btnEncrypt.setGeometry(QtCore.QRect(330, 310, 151, 61))
        font = QtGui.QFont()
        font.setPointSize(10)
        self.btnEncrypt.setFont(font)
        self.btnEncrypt.setObjectName("btnEncrypt")

        self.btnSendtext = QtWidgets.QPushButton(self.centralwidget)
        self.btnSendtext.setGeometry(QtCore.QRect(330, 600, 151, 61))
        font = QtGui.QFont()
        font.setPointSize(10)
        self.btnSendtext.setFont(font)
        self.btnSendtext.setObjectName("btnSendtext")

        self.btnReceivetext = QtWidgets.QPushButton(self.centralwidget)
        self.btnReceivetext.setGeometry(QtCore.QRect(890, 30, 161, 51))
        font = QtGui.QFont()
        font.setPointSize(10)
        self.btnReceivetext.setFont(font)
        self.btnReceivetext.setObjectName("btnReceivetext")

        self.btnDecrypt = QtWidgets.QPushButton(self.centralwidget)
        self.btnDecrypt.setGeometry(QtCore.QRect(900, 310, 151, 61))
        font = QtGui.QFont()
        font.setPointSize(10)
        self.btnDecrypt.setFont(font)
        self.btnDecrypt.setObjectName("btnDecrypt")

        self.btnClearAll = QtWidgets.QPushButton(self.centralwidget)
        self.btnClearAll.setGeometry(QtCore.QRect(900, 600, 151, 61))
        font = QtGui.QFont()
        font.setPointSize(10)
        self.btnClearAll.setFont(font)
        self.btnClearAll.setObjectName("btnClearAll")

        self.label_4 = QtWidgets.QLabel(self.centralwidget)
        self.label_4.setGeometry(QtCore.QRect(610, 300, 31, 20))

        self.btnOpenFile = QtWidgets.QPushButton(self.centralwidget)
        self.btnOpenFile.setGeometry(QtCore.QRect(495, 110, 100, 50))
        font = QtGui.QFont()
        font.setPointSize(10)
        self.btnOpenFile.setFont(font)
        self.btnOpenFile.setObjectName("btnOpenFile")

        self.btnSaveFile = QtWidgets.QPushButton(self.centralwidget)
        self.btnSaveFile.setGeometry(QtCore.QRect(495, 390, 100, 50))
        font = QtGui.QFont()
        font.setPointSize(10)
        self.btnSaveFile.setFont(font)
        self.btnSaveFile.setObjectName("btnSaveFile")

        self.btnGetFile = QtWidgets.QPushButton(self.centralwidget)
        self.btnGetFile.setGeometry(QtCore.QRect(650, 30, 100, 50))
        font = QtGui.QFont()
        font.setPointSize(10)
        self.btnGetFile.setFont(font)
        self.btnGetFile.setObjectName("btnGetFile")

        font = QtGui.QFont()
        font.setPointSize(10)
        self.label_4.setFont(font)
        self.label_4.setObjectName("label_4")

        self.txta = QtWidgets.QLineEdit(self.centralwidget)
        self.txta.setGeometry(QtCore.QRect(650, 300, 71, 22))
        font = QtGui.QFont()
        font.setPointSize(10)
        self.txta.setFont(font)
        self.txta.setObjectName("txta")

        self.label_5 = QtWidgets.QLabel(self.centralwidget)
        self.label_5.setGeometry(QtCore.QRect(610, 330, 31, 20))
        font = QtGui.QFont()
        font.setPointSize(10)
        self.label_5.setFont(font)
        self.label_5.setObjectName("label_5")

        self.txtP2 = QtWidgets.QLineEdit(self.centralwidget)
        self.txtP2.setGeometry(QtCore.QRect(650, 330, 71, 22))
        font = QtGui.QFont()
        font.setPointSize(10)
        self.txtP2.setFont(font)
        self.txtP2.setObjectName("txtP2")

        self.textBox1 = QtWidgets.QTextEdit(self.centralwidget)
        self.textBox1.setAcceptRichText(True)
        self.textBox1.setGeometry(QtCore.QRect(40, 110, 441, 181))
        self.textBox1.setObjectName("textBox1")

        self.textBox3 = QtWidgets.QTextEdit(self.centralwidget)
        self.textBox3.setGeometry(QtCore.QRect(610, 110, 441, 181))
        self.textBox3.setObjectName("textBox3")

        self.textBox4 = QtWidgets.QTextEdit(self.centralwidget)
        self.textBox4.setGeometry(QtCore.QRect(610, 390, 441, 181))
        self.textBox1.setAcceptRichText(True)
        self.textBox4.setReadOnly(True)
        self.textBox4.setObjectName("textBox4")

        self.textBox2 = QtWidgets.QTextEdit(self.centralwidget)
        self.textBox2.setGeometry(QtCore.QRect(40, 390, 441, 181))
        self.textBox2.setReadOnly(True)
        self.textBox2.setTextInteractionFlags(
            QtCore.Qt.TextSelectableByKeyboard | QtCore.Qt.TextSelectableByMouse)
        self.textBox2.setObjectName("textBox2")

        self.txtSavea = QtWidgets.QLineEdit(self.centralwidget)
        self.txtSavea.setGeometry(QtCore.QRect(170, 600, 1, 1))
        self.txtSavea.setMaximumSize(QtCore.QSize(1, 1))
        font = QtGui.QFont()
        font.setPointSize(10)
        self.txtSavea.setFont(font)
        self.txtSavea.setReadOnly(True)
        self.txtSavea.setObjectName("txtSavea")
        self.txtSavek = QtWidgets.QLineEdit(self.centralwidget)
        self.txtSavek.setGeometry(QtCore.QRect(180, 600, 1, 1))
        self.txtSavek.setMaximumSize(QtCore.QSize(1, 1))
        font = QtGui.QFont()
        font.setPointSize(10)
        self.txtSavek.setFont(font)
        self.txtSavek.setReadOnly(True)
        self.txtSavek.setObjectName("txtSavek")
        MainWindow.setCentralWidget(self.centralwidget)
        self.menubar = QtWidgets.QMenuBar(MainWindow)
        self.menubar.setGeometry(QtCore.QRect(0, 0, 1111, 26))
        self.menubar.setObjectName("menubar")
        MainWindow.setMenuBar(self.menubar)
        self.statusbar = QtWidgets.QStatusBar(MainWindow)
        self.statusbar.setObjectName("statusbar")
        MainWindow.setStatusBar(self.statusbar)

        #Hết form

        self.retranslateUi(MainWindow)
        QtCore.QMetaObject.connectSlotsByName(MainWindow)
        self.btnGKey.clicked.connect(self.postKey)
        self.btnEncrypt.clicked.connect(self.encrypt)
        self.btnSendtext.clicked.connect(self.sendCipher)
        self.btnReceivetext.clicked.connect(self.receiveCipher)
        self.btnDecrypt.clicked.connect(self.decrypt)
        self.btnClearAll.clicked.connect(self.clearall)
        self.btnOpenFile.clicked.connect(self.openFile)
        self.btnSaveFile.clicked.connect(self.saveFile)
        self.btnGetFile.clicked.connect(self.getFile)

    checkClick = 0 

    def primesInRange(self, x,y):
        prime_list = []
        for n in range(x, y):
            isPrime = True
            for num in range(2, n):
                if n % num == 0:
                    isPrime = False 
            if isPrime:
                prime_list.append(n)
        return prime_list

    def power(self, a, b, p):
        res = 1     # Khởi tạo kết quả
        y = a
        if (a == 0):
            return 0
        while (b > 0):
            if b % 2 != 0:
                res = (res * y) % p
            b = int(b/2)    
            y = (y * y) % p
        return res % p

    def checkPrimeNumber(self, n):
        if n < 2:
            return False
        for i in range(2, int(n ** (1/2))):
            if n % i == 0:
                return False
        return True

    # tinh UCLN
    def gcd(self, a, b):
        if a < b:
            return self.gcd(b, a)
        elif a % b == 0:
            return b
        else:
            return self.gcd(b, a % b)

    def checkKey(self, MainWindow):
        if self.txtAlpha.text() == "":
            self.msgError("Alpha không được để trống.")
            return False 
        elif self.txtP.text() == "":
            self.msgError("P không được để trống.")
            return False
        elif str(self.txtAlpha.text()).isdigit() == False:
            self.msgError("Alpha phải là số nguyên.")
            return False
        elif str(self.txtP.text()).isdigit() == False:
            self.msgError("P phải là số nguyên.")
            return False
        else:
            return True

    def postKey(self, MainWindow):
        if self.txtAlpha.text() == "" and self.txtP.text()=="":
            primeList = self.primesInRange(8000, 8999)
            P = random.choice(primeList)
            alpha = random.randint(100, 8999)
            self.txtAlpha.setText(str(alpha))
            self.txtP.setText(str(P))

            a = random.randint(1, P-2)
            k = random.randint(1, P-1)
            while self.gcd(P, k) != 1:
                k = random.randint(1, P-1)

            beta = self.power(alpha, a, P)
            self.txtBeta.setText(str(beta))
            self.txtSavea.setText(str(a))
            self.txtSavek.setText(str(k))
        else:
            if self.checkKey(self) == True:
                alpha = int(self.txtAlpha.text())
                P = int(self.txtP.text())
                if (self.checkPrimeNumber(P) == False):
                    self.msgError("P không phải số nguyên tố! Nhập lại.")
                    self.txtP.setText("")
                else:
                    a = random.randint(1, P-2)
                    k = random.randint(1, P-1)
                    while self.gcd(P, k) != 1:
                        k = random.randint(1, P-1)

                    beta = self.power(alpha, a, P)
                    self.txtBeta.setText(str(beta))
                    self.txtSavea.setText(str(a))
                    self.txtSavek.setText(str(k))


    def encrypt(self, MainWindow):
        if self.txtBeta.text() == '':
            self.msgError("Chưa có khóa. Cần tạo khóa.")
        elif self.textBox1.toPlainText() == '':
            self.msgError("Không có bản rõ.")
        else:
            alpha = int(self.txtAlpha.text())
            P = int(self.txtP.text())
            a = int(self.txtSavea.text())
            msg = self.textBox1.toPlainText()
            beta = int(self.txtBeta.text())
            k = int(self.txtSavek.text())
            y1 = self.power(alpha, k, P)
            en_msg = []
            str_msg = ""
            # chuyển đổi chuỗi thành aray
            for i in range(0, len(msg)):
                en_msg.append(msg[i])

            # encrypt array
            for i in range(0, len(en_msg)):
                en_msg[i] = chr(ord(en_msg[i])*pow(beta, k) % P)
                str_msg += en_msg[i]
                #print(en_msg[i])
                #print(ord(en_msg[i]))

            self.textBox2.setText(str_msg)
            return y1, str_msg

    def sendCipher(self, MainWindow):
        if self.txtBeta.text() == '':
            self.msgError("Chưa có khóa. Cần tạo khóa.")
        elif self.textBox2.toPlainText() == '':
            self.msgError("Không có bản mã.")
        else:
            P = int(self.txtP.text())
            a = self.txtSavea.text()
            msg = self.textBox2.toPlainText()
            self.checkClick = 1
            return a, P, msg, self.checkClick

    def receiveCipher(self, MainWindow):
        if self.textBox2.toPlainText() == '':
            self.msgError("Không nhận được thông tin mã hoá.")
        elif self.checkClick == 0:
            self.msgError("Không nhận được thông tin mã hoá.")
        else:
            receive = self.sendCipher(self)
            self.txta.setText(str(receive[0]))
            self.txtP2.setText(str(receive[1]))
            self.textBox3.setText(str(receive[2]))

    def decrypt(self, MainWindow):
        if self.textBox3.toPlainText() == '':
            self.msgError("Không nhận được bản mã.")
        else:
            dr_msg = []
            str_msg = ""
            cipher = self.encrypt(self)
            y1 = cipher[0]
            en_msg = self.textBox3.toPlainText()
            a = int(self.txta.text())
            P = int(self.txtP2.text())
            # Gia ma van ban
            for i in range(0, len(en_msg)):
                dr_msg.append(ord(en_msg[i])*(pow(y1, P-1-a)) % P)
                str_msg += chr(dr_msg[i])
            self.textBox4.setText(str_msg)

    def clearall(self, MainWindow):
        self.txtAlpha.setText("")
        self.txta.setText("")
        self.txtBeta.setText("")
        self.txtP.setText("")
        self.txtP2.setText("")
        self.txtSavea.setText("")
        self.textBox1.setText("")
        self.textBox2.setText("")
        self.textBox3.setText("")
        self.textBox4.setText("")
        self.checkClick = 0

    def get_para_data(self, output_doc_name, paragraph):
        """
        Write the run to the new file and then set its font, bold, alignment, color etc. data.
        """
        output_para = output_doc_name.add_paragraph()
        for run in paragraph.runs:
            output_run = output_para.add_run(run.text)
            # Dữ liệu in đậm
            output_run.bold = run.bold
            # Dữ liệu in nghiêng
            output_run.italic = run.italic
            # Dữ liệu gạch chân
            output_run.underline = run.underline
            # Dữ liệu màu
            output_run.font.color.rgb = run.font.color.rgb
            # Dữ liệu phông chữ
            output_run.style.name = run.style.name
        # Dữ liệu căn chỉnh đooạn
        output_para.paragraph_format.alignment = paragraph.paragraph_format.alignment
        pyperclip.copy(output_para)


    def openFile(self, MainWindow):
        print("Open file")
        fname = QFileDialog.getOpenFileName(self, 'Open file', './/',"Doc files (*.doc *.docx);;Text files (*.txt)")
        print(fname)
        if 'txt' in fname[0]:
            f = open(fname[0], 'r', encoding='utf-8')
            with f:
                data = f.read()
                self.textBox1.setText(data)
        elif 'docx' in fname[0]:
            f = docx.Document(fname[0])
            fullText = []
        
            for para in f.paragraphs:
                for run in para.runs:
                    fullText.append(run.text)
            self.textBox1.setText('\n'.join(fullText))


    def saveFile(self, MainWindow):
        cipher = self.encrypt(self)
        y1 = cipher[0]
        name = QFileDialog.getSaveFileName(self, 'Save File', './/', 'Text files (*.txt);;Doc files (*.docx *.doc)')

        if 'txt' in name[0]:
            file = open(name[0], 'w', encoding='utf-8')
            text = self.textBox2.toPlainText()
            file.write(text)
            file.close
        elif 'doc' in name[0]:
            doc = docx.Document()
            text = self.textBox2.toPlainText()
            doc.add_paragraph(text)
            doc.save(name[0])


    def getFile(self, MainWindow):
        fname = QFileDialog.getOpenFileName(self, 'Open file', './/',"Doc files (*.docx *.doc);;Text files (*.txt)")
        print(fname)
        if 'txt' in fname[0]:
            f = open(fname[0], 'r', encoding='utf-8')

            with f:
                data = f.read()
                self.textBox3.setText(data)
        elif 'docx' in fname[0]:
            f = docx.Document(fname[0])
            fullText = []
            for para in f.paragraphs:
                for p in para.runs:
                    print(para.text)
                    fullText.append(p.text)
            self.textBox3.setText('\n'.join(fullText))

        P = self.txtP.text()
        a = self.txtSavea.text()
        self.checkClick = 1      
        self.txta.setText(a)
        self.txtP2.setText(P)


    def msgError(self, msg):
        msgerr = QtWidgets.QMessageBox()
        msgerr.setIcon(QtWidgets.QMessageBox.Information)
        msgerr.setText("Error")
        msgerr.setInformativeText(msg)
        msgerr.setWindowTitle("Thông báo")
        msgerr.exec_()

    def retranslateUi(self, MainWindow):
        _translate = QtCore.QCoreApplication.translate
        MainWindow.setWindowTitle(_translate("MainWindow", " DuongThiHongVan-Ma hoa Elgamal"))
        self.btnGKey.setText(_translate("MainWindow", "Tạo khóa"))
        self.label.setText(_translate("MainWindow", "Alpha"))
        self.label_2.setText(_translate("MainWindow", "Beta"))
        self.label_3.setText(_translate("MainWindow", "P"))
        self.btnEncrypt.setText(_translate("MainWindow", "Mã hóa"))
        self.btnSendtext.setText(_translate("MainWindow", "Gửi bản mã"))
        self.btnReceivetext.setText(_translate(
            "MainWindow", "Nhận bản mã"))
        self.btnDecrypt.setText(_translate("MainWindow", "Giải mã"))
        self.btnClearAll.setText(_translate("MainWindow", "Xóa tất cả"))
        self.btnOpenFile.setText(_translate("MainWindow", "Open File"))
        self.btnSaveFile.setText(_translate("MainWindow", "Save File"))
        self.btnGetFile.setText(_translate("MainWindow", "Get File"))
        self.label_4.setText(_translate("MainWindow", "a"))
        self.label_5.setText(_translate("MainWindow", "P"))


if __name__ == "__main__":
    import sys
    app = QtWidgets.QApplication(sys.argv)
    MainWindow = QtWidgets.QMainWindow()
    ui = Ui_MainWindow()
    ui.setupUi(MainWindow)
    MainWindow.show()
    sys.exit(app.exec_())

